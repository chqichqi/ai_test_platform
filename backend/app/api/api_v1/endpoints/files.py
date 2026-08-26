"""
文件上传和管理接口
"""

import os
import uuid
from datetime import datetime
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.core.config import settings
from app.core.logger import logger
from app.core.middleware.permission_middleware import Permissions

router = APIRouter()


def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档中提取文本内容，智能转换为规范 Markdown 格式
    
    支持识别：
    - Heading 样式（标准标题）
    - 字体大小/加粗（非标准标题）
    - 中式编号（一、二、三、1. 2. 3. 等）
    """
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document(file_path)
        
        full_text = []
        prev_para_size = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查段落样式，转换为Markdown标题
            style_name = para.style.name.lower()
            
            # 标准样式识别
            if 'heading 1' in style_name or 'title' in style_name:
                full_text.append(f"# {text}")
            elif 'heading 2' in style_name:
                full_text.append(f"## {text}")
            elif 'heading 3' in style_name:
                full_text.append(f"### {text}")
            elif 'heading 4' in style_name:
                full_text.append(f"#### {text}")
            else:
                # 非标准样式：通过字体大小和加粗识别
                is_bold = False
                font_size = None
                
                if para.runs:
                    for run in para.runs:
                        if run.bold:
                            is_bold = True
                        if run.font.size:
                            font_size = run.font.size.pt
                
                # 中式编号识别：一、二、三、1. 2. 等
                chinese_num_pattern = False
                import re
                if re.match(r'^[一二三四五六七八九十]+[、.．]', text):
                    # 中式编号如 "一、登录功能" → 二级标题（功能模块）
                    chinese_num_pattern = True
                    full_text.append(f"## {text}")
                elif re.match(r'^\d+[、.．]\s*[^\d]', text) and len(text) < 30:
                    # 数字编号如 "1. xxx"，短文本 → 三级标题（子功能）
                    chinese_num_pattern = True
                    full_text.append(f"### {text}")
                elif re.match(r'^\d+\.\d+[、.．\s]', text) and len(text) < 40:
                    # 二级数字编号如 "1.1 xxx" → 四级标题
                    chinese_num_pattern = True
                    full_text.append(f"#### {text}")
                elif re.match(r'^第[一二三四五六七八九十\d]+[章节部分]', text):
                    # "第一章"、"第1节"等 → 一级标题（章节）
                    chinese_num_pattern = True
                    full_text.append(f"# {text}")
                
                # 字体大小识别（非中式编号时）
                if not chinese_num_pattern:
                    if font_size and font_size >= 18 and is_bold:
                        full_text.append(f"# {text}")
                    elif font_size and font_size >= 16 and is_bold:
                        full_text.append(f"## {text}")
                    elif font_size and font_size >= 14 and is_bold:
                        full_text.append(f"### {text}")
                    elif is_bold and len(text) < 50 and not text.endswith('.'):
                        # 加粗短文本可能是标题
                        full_text.append(f"### {text}")
                    else:
                        full_text.append(text)
        
        # 提取表格内容
        for table in doc.tables:
            full_text.append("")
            full_text.append("| 表格内容 |")
            full_text.append("|---|")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append("| " + " | ".join(row_text) + " |")
        
        result = "\n".join(full_text)
        
        # 后处理：规范化标题层级
        result = _normalize_markdown_structure(result)
        
        # 统计转换结果
        h1_count = sum(1 for line in result.split('\n') if line.startswith('# '))
        h2_count = sum(1 for line in result.split('\n') if line.startswith('## '))
        h3_count = sum(1 for line in result.split('\n') if line.startswith('### '))
        logger.info(f"Word文档转换完成: 一级标题{h1_count}个, 二级标题{h2_count}个, 三级标题{h3_count}个")
        
        return result
    except Exception as e:
        logger.error(f"解析 Word 文档失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return ""


def _normalize_markdown_structure(text: str) -> str:
    """规范化 Markdown 结构，确保标题层级合理
    
    - 如果没有一级标题，将第一个二级标题提升为一级
    - 确保标题层级连续（不跳级）
    - 添加文档标题（如果缺失）
    """
    lines = text.split('\n')
    
    # 检查是否有一级标题
    has_h1 = any(line.startswith('# ') for line in lines)
    
    # 如果没有一级标题，尝试添加文档标题
    if not has_h1:
        # 查找第一个二级标题，提升为一级
        for i, line in enumerate(lines):
            if line.startswith('## '):
                lines[i] = '# ' + line[3:]
                break
    
    # 确保标题层级连续
    prev_level = 0
    for i, line in enumerate(lines):
        if line.startswith('#'):
            # 计算当前标题层级
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break
            
            # 如果跳级（如从 # 直接跳到 ###），调整
            if prev_level > 0 and level > prev_level + 1:
                # 调整为上一级+1
                new_level = prev_level + 1
                lines[i] = '#' * new_level + line[level:]
                level = new_level
            
            prev_level = level
    
    return '\n'.join(lines)


def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 文档中提取文本内容"""
    try:
        import fitz
        doc = fitz.open(file_path)
        full_text = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                full_text.append(text)
        doc.close()
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"解析 PDF 文档失败: {str(e)}")
        return ""


def ensure_upload_dir():
    """确保上传目录存在"""
    upload_dir = os.path.join(settings.UPLOAD_DIR, "requirement_docs")
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir, exist_ok=True)
    return upload_dir


@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: dict = Depends(Permissions.VERSION_CREATE)
):
    """
    上传需求文档文件
    
    支持格式：docx, doc, pdf, md, txt
    返回文件路径和类型
    """
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文件名不能为空"
        )
    
    # 获取文件扩展名
    file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
    
    # 验证文件类型
    allowed_extensions = ['docx', 'doc', 'pdf', 'md', 'txt', 'markdown']
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件类型：{file_ext}。支持的类型：{', '.join(allowed_extensions)}"
        )
    
    # 验证文件大小
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    if file_size > settings.MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"文件大小超过限制（最大 {settings.MAX_UPLOAD_SIZE // 1024 // 1024}MB）"
        )
    
    # 生成唯一文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex[:8]
    safe_filename = f"{timestamp}_{unique_id}_{file.filename}"
    
    # 确保上传目录存在
    upload_dir = ensure_upload_dir()
    
    # 保存文件
    file_path = os.path.join(upload_dir, safe_filename)
    
    try:
        with open(file_path, "wb") as f:
            content = file.file.read()
            f.write(content)
        
        logger.info(f"文件上传成功：{file.filename} -> {file_path}")
        
        relative_path = f"requirement_docs/{safe_filename}"
        
        extracted_text = ""
        ocr_images = []
        
        # 使用统一的文档解析服务
        from app.core.services.document_parser import document_parser
        
        if file_ext in ['docx', 'doc', 'pdf', 'md', 'txt', 'markdown']:
            result = document_parser.parse_file(file_path, file.filename)
            extracted_text = result.get('content', '')
            ocr_images = result.get('images', [])
            logger.info(f"文档解析完成，提取 {len(extracted_text)} 字符，图片OCR {len(ocr_images)} 张")
        
        return {
            "success": True,
            "file_path": relative_path,
            "file_type": file_ext,
            "file_name": file.filename,
            "file_size": file_size,
            "extracted_text": extracted_text,
            "ocr_images_count": len(ocr_images),
            "message": "文件上传成功"
        }
        
    except Exception as e:
        logger.error(f"文件上传失败：{str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文件上传失败：{str(e)}"
        )


@router.get("/download/{file_path:path}")
async def download_file(
    file_path: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(Permissions.VERSION_READ)
):
    """
    下载需求文档文件
    """
    # 构建完整路径
    full_path = os.path.join(settings.UPLOAD_DIR, file_path)
    
    if not os.path.exists(full_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文件不存在"
        )
    
    # 获取文件名
    filename = os.path.basename(full_path)
    
    return FileResponse(
        path=full_path,
        filename=filename,
        media_type="application/octet-stream"
    )


@router.get("/preview/{file_path:path}")
async def preview_file(
    file_path: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(Permissions.VERSION_READ)
):
    """
    预览需求文档文件
    
    对于 PDF：直接返回文件
    对于 Office 文档：返回文件下载链接（可配合 Office Online 预览）
    对于文本文件：返回文件内容
    """
    full_path = os.path.join(settings.UPLOAD_DIR, file_path)
    
    if not os.path.exists(full_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文件不存在"
        )
    
    file_ext = file_path.split('.')[-1].lower()
    filename = os.path.basename(full_path)
    
    # PDF 文件直接返回
    if file_ext == 'pdf':
        return FileResponse(
            path=full_path,
            filename=filename,
            media_type="application/pdf"
        )
    
    # Office 文档返回文件（前端可使用 Office Online 预览）
    elif file_ext in ['docx', 'doc']:
        return FileResponse(
            path=full_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # 文本文件返回内容
    elif file_ext in ['md', 'txt', 'markdown']:
        try:
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {
                "success": True,
                "content": content,
                "file_type": file_ext
            }
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"读取文件失败：{str(e)}"
            )
    
    else:
        return FileResponse(
            path=full_path,
            filename=filename,
            media_type="application/octet-stream"
        )


@router.post("/analyze", status_code=status.HTTP_200_OK)
async def analyze_document(
    request_body: dict = {},
    db: Session = Depends(get_db),
    current_user: dict = Depends(Permissions.VERSION_CREATE)
):
    """
    使用 LLM 智能分析文档内容，生成标准功能模块格式
    
    Request Body:
        - content: 文档内容（可选）
        - file_path: 文件路径（可选，与 content 二选一）
        - document_type: 文档类型（docx/pdf/md/txt）
    
    Returns:
        分析结果，包含功能模块列表和转换后的 Markdown 内容
    """
    from app.core.services.doc_preprocess_service import DocumentPreprocessService
    
    # 从 request body 获取参数
    content = request_body.get("content", "")
    file_path = request_body.get("file_path")
    document_type = request_body.get("document_type", "unknown")
    
    # 获取文档内容
    doc_content = content or ""
    
    if file_path and not content:
        # 从文件读取内容
        full_path = os.path.join(settings.UPLOAD_DIR, file_path)
        if os.path.exists(full_path):
            file_ext = file_path.split('.')[-1].lower()
            if file_ext in ['docx', 'doc']:
                doc_content = extract_text_from_docx(full_path)
            elif file_ext == 'pdf':
                doc_content = extract_text_from_pdf(full_path)
            elif file_ext in ['md', 'txt', 'markdown']:
                with open(full_path, 'r', encoding='utf-8') as f:
                    doc_content = f.read()
            logger.info(f"从文件读取内容，类型: {file_ext}, 长度: {len(doc_content)}")
    
    if not doc_content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="请提供文档内容或有效的文件路径"
        )
    
    # 调用 LLM 分析
    service = DocumentPreprocessService(db)
    result = await service.analyze_document(doc_content, document_type)
    
    if not result.get("success"):
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=result.get("error", "文档分析失败")
        )
    
    return {
        "success": True,
        "document_title": result.get("document_title"),
        "modules": result.get("modules", []),
        "markdown_content": result.get("markdown_content"),
        "stats": result.get("stats"),
        "message": "文档分析完成"
    }