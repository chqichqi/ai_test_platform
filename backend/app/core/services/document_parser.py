"""
文档解析服务
支持 Word、PDF、Markdown、文本等格式解析
集成 OCR 功能处理文档中的图片
支持文档格式规范化处理（标题识别、中式编号转换）
"""

import os
import tempfile
from typing import Optional, Dict, List, Any
from pathlib import Path

from app.core.logger import logger


class DocumentParserService:
    """文档解析服务"""
    
    def __init__(self):
        self.supported_formats = {
            '.docx': 'word',
            '.doc': 'word',
            '.pdf': 'pdf',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.txt': 'text',
        }
    
    def parse_file(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        解析文件
        
        Args:
            file_path: 文件路径
            file_name: 文件名称
            
        Returns:
            解析结果：{
                "content": "文本内容（已规范化为Markdown格式）",
                "format": "文档格式",
                "images": ["图片 OCR 文本 1", "图片 OCR 文本 2"],
                "metadata": {"pages": 10, "words": 1000}
            }
        """
        file_ext = Path(file_name).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式：{file_ext}")
        
        format_type = self.supported_formats[file_ext]
        
        try:
            if format_type == 'word':
                return self._parse_word(file_path)
            elif format_type == 'pdf':
                return self._parse_pdf(file_path)
            elif format_type == 'markdown':
                return self._parse_markdown(file_path)
            elif format_type == 'text':
                return self._parse_text(file_path)
            else:
                raise ValueError(f"未知的文档格式：{format_type}")
        except Exception as e:
            logger.error(f"解析文件失败：{str(e)}")
            raise
    
    def _parse_word(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档 - 使用格式规范化处理
        
        支持：
        - Heading 样式识别（标准标题）
        - 字体大小/加粗识别（非标准标题）
        - 中式编号识别（一、二、三、1. 2. 3. 等）
        - 表格提取
        - 图片 OCR
        """
        try:
            from docx import Document
            from docx.shared import Pt
            import re
            
            doc = Document(file_path)
            full_text = []
            images = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                style_name = para.style.name.lower()
                
                # 标准样式识别
                if 'heading 1' in style_name or 'title' in style_name:
                    full_text.append(f"# {text}")
                elif 'heading 2' in style_name:
                    full_text.append(f"## {text}")
                elif 'heading 3' in style_name:
                    full_text.append(f"### {text}")
                elif 'heading 4' in style_name:
                    full_text.append(f"#### {text}")
                else:
                    # 非标准样式：通过字体大小和加粗识别
                    is_bold = False
                    font_size = None
                    
                    if para.runs:
                        for run in para.runs:
                            if run.bold:
                                is_bold = True
                            if run.font.size:
                                font_size = run.font.size.pt
                    
                    # 中式编号识别
                    chinese_num_pattern = False
                    if re.match(r'^[一二三四五六七八九十]+[、.．]', text):
                        chinese_num_pattern = True
                        full_text.append(f"## {text}")
                    elif re.match(r'^\d+[、.．]\s*[^\d]', text) and len(text) < 30:
                        chinese_num_pattern = True
                        full_text.append(f"### {text}")
                    elif re.match(r'^\d+\.\d+[、.．\s]', text) and len(text) < 40:
                        chinese_num_pattern = True
                        full_text.append(f"#### {text}")
                    elif re.match(r'^第[一二三四五六七八九十\d]+[章节部分]', text):
                        chinese_num_pattern = True
                        full_text.append(f"# {text}")
                    
                    # 字体大小识别（非中式编号时）
                    if not chinese_num_pattern:
                        if font_size and font_size >= 18 and is_bold:
                            full_text.append(f"# {text}")
                        elif font_size and font_size >= 16 and is_bold:
                            full_text.append(f"## {text}")
                        elif font_size and font_size >= 14 and is_bold:
                            full_text.append(f"### {text}")
                        elif is_bold and len(text) < 50 and not text.endswith('.'):
                            full_text.append(f"### {text}")
                        else:
                            full_text.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                full_text.append("")
                full_text.append("| 表格内容 |")
                full_text.append("|---|")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append("| " + " | ".join(row_text) + " |")
            
            # 提取图片并 OCR
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        ocr_text = self._ocr_image(rel.target_part.blob)
                        if ocr_text:
                            images.append(ocr_text)
                            full_text.append(f"\n### 图片内容\n{ocr_text}\n")
                    except Exception as e:
                        logger.warning(f"图片 OCR 失败：{str(e)}")
            
            content = "\n".join(full_text)
            
            # 规范化 Markdown 结构
            content = self._normalize_markdown_structure(content)
            
            h1_count = sum(1 for line in content.split('\n') if line.startswith('# '))
            h2_count = sum(1 for line in content.split('\n') if line.startswith('## '))
            h3_count = sum(1 for line in content.split('\n') if line.startswith('### '))
            logger.info(f"Word文档转换完成: 一级标题{h1_count}个, 二级标题{h2_count}个, 三级标题{h3_count}个, 图片OCR{len(images)}张")
            
            return {
                "content": content,
                "format": "word",
                "images": images,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "images": len(images),
                    "h1_count": h1_count,
                    "h2_count": h2_count,
                    "h3_count": h3_count
                }
            }
        except ImportError:
            logger.warning("python-docx 未安装，使用简单文本解析")
            return self._parse_text(file_path)
        except Exception as e:
            logger.error(f"解析 Word 文档失败：{str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文档"""
        try:
            import fitz
            
            doc = fitz.open(file_path)
            paragraphs = []
            images = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text and text.strip():
                    paragraphs.append(f"## 第{page_num + 1}页\n\n{text}")
                
                for img_index, img in enumerate(page.get_images(full=True)):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        ocr_text = self._ocr_image(image_bytes)
                        if ocr_text:
                            images.append(ocr_text)
                            paragraphs.append(f"\n### 图片内容（第{page_num + 1}页）\n{ocr_text}\n")
                    except Exception as e:
                        logger.warning(f"PDF 图片 OCR 失败：{str(e)}")
            
            doc.close()
            content = "\n\n".join(paragraphs)
            
            return {
                "content": content,
                "format": "pdf",
                "images": images,
                "metadata": {
                    "pages": len(paragraphs),
                    "images": len(images)
                }
            }
        except ImportError:
            logger.warning("fitz(PyMuPDF) 未安装，尝试 pdfplumber")
            return self._parse_pdf_with_plumber(file_path)
        except Exception as e:
            logger.error(f"解析 PDF 文档失败：{str(e)}")
            raise
    
    def _parse_pdf_with_plumber(self, file_path: str) -> Dict[str, Any]:
        """使用 pdfplumber 解析 PDF"""
        try:
            import pdfplumber
            
            paragraphs = []
            images = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        paragraphs.append(f"## 第{page_num + 1}页\n\n{text}")
                    
                    if page.images:
                        images.append(f"[图片 - 第{page_num + 1}页]")
            
            content = "\n\n".join(paragraphs)
            
            return {
                "content": content,
                "format": "pdf",
                "images": images,
                "metadata": {
                    "pages": len(pdf.pages),
                    "images": len(images)
                }
            }
        except Exception as e:
            logger.error(f"pdfplumber 解析失败：{str(e)}")
            raise ValueError("无法解析 PDF 文件，请安装 PyMuPDF 或 pdfplumber")
    
    def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """解析 Markdown 文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                "content": content,
                "format": "markdown",
                "images": [],
                "metadata": {
                    "lines": len(content.split('\n')),
                    "characters": len(content)
                }
            }
        except Exception as e:
            logger.error(f"解析 Markdown 文档失败：{str(e)}")
            raise
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析文本文档"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                raise ValueError("无法识别文件编码")
            
            return {
                "content": content,
                "format": "text",
                "images": [],
                "metadata": {
                    "lines": len(content.split('\n')),
                    "characters": len(content)
                }
            }
        except Exception as e:
            logger.error(f"解析文本文档失败：{str(e)}")
            raise
    
    def _ocr_image(self, image_data: bytes) -> Optional[str]:
        """
        对图片进行 OCR 识别
        
        Args:
            image_data: 图片二进制数据
            
        Returns:
            OCR 识别的文本内容
        """
        try:
            from app.core.services.ocr_service import OCRService
            
            ocr_service = OCRService()
            result = ocr_service.recognize_text(image_data, language='chi_sim+eng')
            
            if result.get('success') and result.get('text'):
                return result['text'].strip()
            return ""
        except ImportError:
            logger.warning("OCR 服务不可用")
            return ""
        except Exception as e:
            logger.warning(f"图片 OCR 失败：{str(e)}")
            return ""
    
    def _normalize_markdown_structure(self, text: str) -> str:
        """规范化 Markdown 结构，确保标题层级合理
        
        - 如果没有一级标题，将第一个二级标题提升为一级
        - 确保标题层级连续（不跳级）
        """
        lines = text.split('\n')
        
        has_h1 = any(line.startswith('# ') for line in lines)
        
        if not has_h1:
            for i, line in enumerate(lines):
                if line.startswith('## '):
                    lines[i] = '# ' + line[3:]
                    break
        
        prev_level = 0
        for i, line in enumerate(lines):
            if line.startswith('#'):
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                if prev_level > 0 and level > prev_level + 1:
                    new_level = prev_level + 1
                    lines[i] = '#' * new_level + line[level:]
                    level = new_level
                
                prev_level = level
        
        return '\n'.join(lines)
    
    def merge_content(self, text_content: str, image_texts: List[str]) -> str:
        """
        合并文本内容和图片 OCR 文本
        
        Args:
            text_content: 文档文本内容
            image_texts: 图片 OCR 文本列表
            
        Returns:
            合并后的完整内容
        """
        if not image_texts:
            return text_content
        
        # 将图片 OCR 文本附加到文档末尾
        image_section = "\n\n---\n\n## 文档图片内容识别\n\n"
        for i, img_text in enumerate(image_texts, 1):
            image_section += f"\n### 图片{i}\n{img_text}\n"
        
        return text_content + image_section


# 全局服务实例
document_parser = DocumentParserService()
